from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('SECTION 2 - PART 2: PAYSTACK API INTEGRATION QUESTIONS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add project information
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Project: ').bold = True
p.add_run('Paystack Payment Integration Platform')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Technology Stack: ').bold = True
p.add_run('FastAPI (Backend & Frontend), Python, Bootstrap 5, Paystack APIs')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Deployed URL: ').bold = True
p.add_run('[Your deployment URL will go here]')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('GitHub Repository: ').bold = True
p.add_run('[Your GitHub URL will go here]')

doc.add_page_break()

# Question 1
doc.add_heading('Question 1: Did you encounter any issues when using the Paystack APIs and how did you resolve them?', level=1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Answer:').bold = True
doc.add_paragraph()

doc.add_paragraph(
    'Yes, I encountered several issues during the integration process. Here\'s a detailed breakdown of each issue '
    'and how I resolved it:'
)

# Issue 1
doc.add_heading('1. Authentication Errors (401 Unauthorized)', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'Initially, I received 401 Unauthorized errors when making API requests to Paystack endpoints. The API was '
    'rejecting my requests despite having valid API keys.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'The issue was twofold: First, I was not formatting the Authorization header correctly. Second, I was initially '
    'using the public key instead of the secret key for server-side requests.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_auth = [
    'Verified the correct Authorization header format: "Bearer {secret_key}"',
    'Ensured I was using the secret key (sk_test_...) for backend API calls, not the public key',
    'Implemented proper environment variable loading using python-dotenv',
    'Created a helper function to generate headers consistently across all API calls',
    'Added validation to check if API keys are loaded correctly on application startup',
]
for step in resolution_steps_auth:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Code Implementation:').italic = True
doc.add_paragraph(
    'def get_paystack_headers():\n'
    '    return {\n'
    '        "Authorization": f"Bearer {PAYSTACK_SECRET_KEY}",\n'
    '        "Content-Type": "application/json"\n'
    '    }',
    style='Normal'
)

# Issue 2
doc.add_heading('2. Amount Format Confusion', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'Payments were being initialized with incorrect amounts. For example, when I tried to charge ₦1000, the system '
    'was attempting to charge ₦10 instead, causing confusion and failed transactions.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'Paystack expects amounts in the smallest currency unit (kobo for NGN, pesewas for GHS, cents for ZAR). '
    'I was initially sending amounts in Naira without conversion.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_amount = [
    'Learned that 1 Naira = 100 Kobo',
    'Implemented automatic conversion: amount_in_kobo = amount * 100',
    'Added conversion back when displaying amounts: amount_in_naira = amount / 100',
    'Created helper functions for currency conversion to maintain consistency',
    'Added validation to ensure amounts are positive and within acceptable ranges',
    'Updated UI to clearly indicate currency (NGN)',
]
for step in resolution_steps_amount:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Code Implementation:').italic = True
doc.add_paragraph(
    '# Convert to kobo for Paystack\n'
    'amount_in_kobo = int(amount * 100)\n\n'
    '# Convert back for display\n'
    'amount_in_naira = transaction_data["amount"] / 100',
    style='Normal'
)

# Issue 3
doc.add_heading('3. Webhook Signature Verification', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'When setting up webhook endpoints, I couldn\'t verify that the webhook requests were actually coming from '
    'Paystack. This posed a security risk as malicious actors could potentially send fake payment confirmations.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'I didn\'t initially implement webhook signature verification, which is crucial for security. Paystack sends '
    'a signature in the x-paystack-signature header that must be verified.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_webhook = [
    'Studied Paystack\'s webhook security documentation',
    'Implemented HMAC SHA512 signature verification',
    'Retrieved the signature from the x-paystack-signature header',
    'Computed expected signature using the secret key and raw request body',
    'Compared computed signature with received signature',
    'Rejected webhooks with invalid signatures to prevent fraud',
    'Tested webhook delivery using Paystack\'s webhook testing tool',
]
for step in resolution_steps_webhook:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Code Implementation:').italic = True
doc.add_paragraph(
    'import hmac\n'
    'import hashlib\n\n'
    'signature = request.headers.get("x-paystack-signature")\n'
    'body = await request.body()\n\n'
    'hash_value = hmac.new(\n'
    '    PAYSTACK_SECRET_KEY.encode(\'utf-8\'),\n'
    '    body,\n'
    '    hashlib.sha512\n'
    ').hexdigest()\n\n'
    'if hash_value != signature:\n'
    '    raise HTTPException(status_code=400, detail="Invalid signature")',
    style='Normal'
)

# Issue 4
doc.add_heading('4. CORS Issues During Development', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'During local development, browsers were blocking requests due to Cross-Origin Resource Sharing (CORS) policies. '
    'This prevented the frontend from communicating with the backend API.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'Modern browsers enforce CORS policies to prevent unauthorized cross-origin requests. Since the frontend and '
    'backend were on the same server but making external API calls, CORS headers needed to be properly configured.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_cors = [
    'Used Paystack\'s official inline.js popup library which handles CORS automatically',
    'The popup solution avoids CORS issues by handling payment in Paystack\'s domain',
    'For direct API calls during testing, used Postman which doesn\'t enforce CORS',
    'Considered adding CORS middleware for production if needed',
    'Ensured all sensitive API calls happen server-side to avoid exposing secret keys',
]
for step in resolution_steps_cors:
    doc.add_paragraph(step, style='List Bullet')

# Issue 5
doc.add_heading('5. Transaction Reference Management', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'Initially, I was manually generating transaction references, which sometimes resulted in duplicate references '
    'or references that didn\'t meet Paystack\'s requirements (only alphanumeric, dash, period, and equals sign allowed).'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'Transaction references must be unique across all transactions. My manual generation method occasionally '
    'created duplicates, especially during testing.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_ref = [
    'Let Paystack auto-generate unique references by not specifying a reference in the initialize request',
    'Paystack generates unique references automatically that are guaranteed to be unique',
    'Stored the returned reference in local storage for transaction tracking',
    'Implemented reference validation for cases where custom references are needed',
    'Added database constraints to prevent duplicate references if using custom generation',
]
for step in resolution_steps_ref:
    doc.add_paragraph(step, style='List Bullet')

# Issue 6
doc.add_heading('6. Callback URL Configuration', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'After successful payment, users weren\'t being redirected back to my application properly. The callback URL '
    'wasn\'t working as expected during local development.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'During local development, localhost URLs aren\'t accessible to Paystack\'s servers. Additionally, I wasn\'t '
    'properly handling the callback URL with the transaction reference parameter.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_callback = [
    'Set up a callback route in the application to handle post-payment redirects',
    'Configured the callback URL in the transaction initialization request',
    'For local testing, used ngrok to create a public URL',
    'Implemented proper parameter handling in the callback route to extract transaction reference',
    'Created a beautiful callback page with automatic verification',
    'For production, used the actual domain name in the callback URL',
]
for step in resolution_steps_callback:
    doc.add_paragraph(step, style='List Bullet')

# Issue 7
doc.add_heading('7. Environment Variable Management', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
doc.add_paragraph(
    'Managing API keys securely across different environments (development, testing, production) was challenging. '
    'I initially hardcoded keys, which is a security risk.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Root Cause: ').bold = True
doc.add_paragraph(
    'Hardcoding API keys in source code exposes them in version control and makes it difficult to switch between '
    'test and live keys for different environments.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
resolution_steps_env = [
    'Implemented python-dotenv for environment variable management',
    'Created .env file for local development (added to .gitignore)',
    'Created .env.example as a template for other developers',
    'Used environment variables for all sensitive configuration',
    'Set up environment variables in deployment platforms (Heroku, Vercel, etc.)',
    'Implemented validation to ensure required environment variables are present',
    'Documented the required environment variables in README',
]
for step in resolution_steps_env:
    doc.add_paragraph(step, style='List Bullet')

# Learning Outcomes
doc.add_page_break()
doc.add_heading('Key Learning Outcomes', level=2)
doc.add_paragraph()

learnings = [
    'Always read API documentation thoroughly before implementation',
    'Test with small amounts first to avoid costly mistakes',
    'Implement proper error handling and logging from the start',
    'Security should never be an afterthought - implement signature verification immediately',
    'Use environment variables for all configuration',
    'Understand currency conversion requirements for payment gateways',
    'Test webhooks using the provider\'s testing tools',
    'Keep detailed logs of API requests and responses during development',
]
for learning in learnings:
    doc.add_paragraph(learning, style='List Bullet')

# Question 2
doc.add_page_break()
doc.add_heading('Question 2: How can you confirm if a transaction was successful or not?', level=1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Answer:').bold = True
doc.add_paragraph()

doc.add_paragraph(
    'There are multiple methods to confirm transaction success with Paystack, each serving different use cases. '
    'A robust implementation should use a combination of these methods for reliability and security.'
)

# Method 1
doc.add_heading('Method 1: Transaction Verification API (Primary & Most Reliable)', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description: ').bold = True
doc.add_paragraph(
    'The Verify Transaction endpoint is the most reliable method to confirm payment status. It queries Paystack\'s '
    'database directly using the transaction reference.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('API Endpoint: ').bold = True
p.add_run('GET /transaction/verify/:reference')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('How It Works: ').bold = True
verify_steps = [
    'After payment, obtain the transaction reference',
    'Make a GET request to the verification endpoint',
    'Receive detailed transaction information including status',
    'Check if status is "success" for successful payments',
    'Extract additional details like amount, currency, channel, and customer info',
]
for step in verify_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Response Data Includes:').bold = True
response_data = [
    'status: Transaction status (success, failed, abandoned, pending)',
    'amount: Amount paid (in kobo)',
    'currency: Currency used (NGN, GHS, ZAR, USD)',
    'channel: Payment method (card, bank, ussd, mobile_money)',
    'paid_at: Timestamp when payment was completed',
    'customer: Customer email and details',
    'authorization: Card authorization details for future charges',
    'metadata: Custom data attached to the transaction',
]
for data in response_data:
    doc.add_paragraph(data, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Implementation Example:').italic = True
doc.add_paragraph(
    'async def verify_payment(reference: str):\n'
    '    response = requests.get(\n'
    '        f"{PAYSTACK_BASE_URL}/transaction/verify/{reference}",\n'
    '        headers={"Authorization": f"Bearer {SECRET_KEY}"}\n'
    '    )\n'
    '    data = response.json()\n'
    '    \n'
    '    if data["status"] and data["data"]["status"] == "success":\n'
    '        return True, data["data"]\n'
    '    return False, None',
    style='Normal'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('When to Use:').bold = True
use_cases_verify = [
    'After payment callback to confirm transaction',
    'When customer claims payment but order isn\'t updated',
    'For reconciliation and accounting purposes',
    'Before delivering goods or services',
    'When webhook delivery fails or is delayed',
]
for use_case in use_cases_verify:
    doc.add_paragraph(use_case, style='List Bullet')

# Method 2
doc.add_heading('Method 2: Webhooks (Recommended for Production)', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description: ').bold = True
doc.add_paragraph(
    'Webhooks provide real-time notifications when transaction events occur. Paystack sends HTTP POST requests '
    'to your configured webhook URL immediately when payment status changes.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Webhook Events: ').bold = True
webhook_events = [
    'charge.success: Payment completed successfully',
    'charge.failed: Payment attempt failed',
    'transfer.success: Transfer completed',
    'transfer.failed: Transfer failed',
    'subscription.create: New subscription created',
    'subscription.disable: Subscription disabled',
]
for event in webhook_events:
    doc.add_paragraph(event, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('How It Works:').bold = True
webhook_steps = [
    'Configure webhook URL in Paystack dashboard (e.g., https://yourapp.com/webhook/paystack)',
    'When payment succeeds, Paystack sends POST request with event data',
    'Verify webhook signature for security (HMAC SHA512)',
    'Extract transaction details from webhook payload',
    'Update database with transaction status',
    'Send confirmation email or SMS to customer',
    'Return 200 OK response to acknowledge receipt',
]
for step in webhook_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Advantages:').bold = True
webhook_advantages = [
    'Real-time: Instant notification when payment succeeds',
    'Efficient: No polling required',
    'Reliable: Paystack retries failed webhooks automatically',
    'Complete Data: Receives full transaction details',
    'Scalable: Works for high-volume transactions',
]
for adv in webhook_advantages:
    doc.add_paragraph(adv, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Security Implementation:').italic = True
doc.add_paragraph(
    'def verify_webhook_signature(payload, signature, secret):\n'
    '    hash_value = hmac.new(\n'
    '        secret.encode(\'utf-8\'),\n'
    '        payload,\n'
    '        hashlib.sha512\n'
    '    ).hexdigest()\n'
    '    return hash_value == signature\n\n'
    '@app.post("/webhook/paystack")\n'
    'async def paystack_webhook(request: Request):\n'
    '    signature = request.headers.get("x-paystack-signature")\n'
    '    body = await request.body()\n'
    '    \n'
    '    if not verify_webhook_signature(body, signature, SECRET_KEY):\n'
    '        raise HTTPException(status_code=400)\n'
    '    \n'
    '    event = await request.json()\n'
    '    if event["event"] == "charge.success":\n'
    '        update_transaction_status(event["data"]["reference"], "success")\n'
    '    \n'
    '    return {"status": "success"}',
    style='Normal'
)

# Method 3
doc.add_heading('Method 3: Payment Callback Response', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description: ').bold = True
doc.add_paragraph(
    'After payment, Paystack redirects the customer to your callback URL with the transaction reference. '
    'This method relies on the redirect but should always be verified using Method 1.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('How It Works:').bold = True
callback_steps = [
    'User completes payment in Paystack popup/page',
    'Paystack redirects to: your-callback-url?reference=TRANSACTION_REF',
    'Extract reference from URL parameters',
    'Call Verify Transaction API to confirm status',
    'Display success or failure message to user',
]
for step in callback_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Important Note:').bold = True
doc.add_paragraph(
    'Never trust the callback alone! Always verify the transaction using the API, as users can manipulate URLs. '
    'The callback is for user experience only - verification is for security.'
)

# Method 4
doc.add_heading('Method 4: List Transactions API', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description: ').bold = True
doc.add_paragraph(
    'Fetches a paginated list of all transactions. Useful for reconciliation, reporting, and bulk status checks.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('API Endpoint: ').bold = True
p.add_run('GET /transaction')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Use Cases:').bold = True
list_use_cases = [
    'Daily reconciliation of payments',
    'Generating financial reports',
    'Checking status of multiple transactions',
    'Audit trails and compliance',
    'Displaying transaction history to users',
]
for use_case in list_use_cases:
    doc.add_paragraph(use_case, style='List Bullet')

# Best Practices
doc.add_page_break()
doc.add_heading('Best Practices for Transaction Verification', level=2)
doc.add_paragraph()

best_practices = [
    'Multiple Verification: Use both webhooks and API verification for reliability',
    
    'Idempotency: Handle duplicate verifications gracefully - verify once, update once',
    
    'Status Mapping: Understand all possible statuses:\n'
    '  • success: Payment completed\n'
    '  • failed: Payment failed\n'
    '  • abandoned: User canceled\n'
    '  • pending: Still processing\n'
    '  • queued: In queue for processing',
    
    'Amount Verification: Always verify the amount matches your expected amount',
    
    'Error Handling: Implement retries for network failures',
    
    'Logging: Log all verification attempts for debugging and audit',
    
    'Timeout Handling: Don\'t wait forever - set reasonable timeouts',
    
    'User Communication: Keep users informed of payment status',
    
    'Database Updates: Update order status immediately after successful verification',
    
    'Security: Never expose secret keys in client-side code',
]
for practice in best_practices:
    doc.add_paragraph(practice, style='List Bullet')

# Complete Flow
doc.add_heading('Complete Verification Flow (Recommended)', level=2)
doc.add_paragraph()

complete_flow = [
    'Customer initiates payment → Initialize Transaction API',
    'Customer completes payment in Paystack → Payment processed',
    'Webhook received → Verify signature, update status to "processing"',
    'Callback redirect → User sees "verifying payment" message',
    'Call Verify Transaction API → Confirm final status',
    'If success: Update database, send confirmation email, show success page',
    'If failed: Update database, notify customer, offer retry option',
    'Background job: Reconcile with List Transactions API daily',
]
for i, step in enumerate(complete_flow, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('This multi-layered approach ensures:').bold = True
ensures = [
    'No missed payments',
    'Real-time customer feedback',
    'Accurate financial records',
    'Protection against fraud',
    'Excellent user experience',
]
for item in ensures:
    doc.add_paragraph(item, style='List Bullet')

# Question 3
doc.add_page_break()
doc.add_heading('Question 3: What are webhooks?', level=1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Answer:').bold = True
doc.add_paragraph()

# Definition
doc.add_heading('Definition and Concept', level=2)
doc.add_paragraph(
    'Webhooks are automated HTTP callbacks that enable real-time, event-driven communication between applications. '
    'When a specific event occurs in one system (the source), it automatically sends data to a predefined URL '
    'endpoint in another system (the destination).'
)
doc.add_paragraph()
doc.add_paragraph(
    'Think of webhooks as a "push" notification system for APIs - instead of your application repeatedly asking '
    '"Has anything changed?" (polling), the source system says "I\'ll tell you when something changes" (webhooks).'
)

# Real-world Analogy
doc.add_heading('Real-World Analogy', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Webhooks are like:').bold = True
doc.add_paragraph()

analogies = [
    'Door bell: Instead of standing outside constantly knocking to see if someone is home, you ring the doorbell '
    'and wait for them to answer.',
    
    'Text message notifications: Your phone alerts you when you receive a message, rather than you having to '
    'open the app every minute to check.',
    
    'Food delivery updates: The restaurant notifies you when your order is ready, rather than you calling them '
    'every 5 minutes to ask.',
]
for analogy in analogies:
    doc.add_paragraph(analogy, style='List Bullet')

# How Webhooks Work
doc.add_heading('How Webhooks Work: Step-by-Step', level=2)
doc.add_paragraph()

webhook_process = [
    'Setup Phase:\n'
    '  • You configure a webhook URL in the source system (e.g., https://yourapp.com/webhooks/paystack)\n'
    '  • You select which events you want to receive notifications for\n'
    '  • You implement an endpoint in your application to receive webhook data',
    
    'Event Occurs:\n'
    '  • Something significant happens in the source system\n'
    '  • Example: A customer completes a payment on Paystack',
    
    'Webhook Triggered:\n'
    '  • The source system detects the event\n'
    '  • Prepares a payload with event details\n'
    '  • Signs the payload for security',
    
    'HTTP Request Sent:\n'
    '  • Source sends POST request to your webhook URL\n'
    '  • Includes event data in JSON format\n'
    '  • Adds security signature in headers',
    
    'Your Server Receives:\n'
    '  • Your webhook endpoint receives the request\n'
    '  • Verifies the signature to ensure authenticity\n'
    '  • Parses the event data',
    
    'Processing:\n'
    '  • Your application processes the event\n'
    '  • Updates database, sends emails, triggers workflows, etc.\n'
    '  • Returns HTTP 200 OK to acknowledge receipt',
    
    'Retry Mechanism:\n'
    '  • If your server doesn\'t respond or returns an error\n'
    '  • Source system automatically retries delivery\n'
    '  • Usually with exponential backoff (e.g., 1 min, 5 min, 15 min)',
]
for i, step in enumerate(webhook_process, 1):
    doc.add_paragraph(f'{i}. {step}')

# Webhooks vs Polling
doc.add_page_break()
doc.add_heading('Webhooks vs Polling: Comparison', level=2)
doc.add_paragraph()

# Create a table-like comparison
doc.add_paragraph().add_run('WEBHOOKS (Push Model)').bold = True
webhook_pros = [
    'Real-time: Instant notifications when events occur',
    'Efficient: Only sends data when something happens',
    'Scalable: No constant API calls needed',
    'Cost-effective: Reduces API usage and server load',
    'Automatic: No need to check for updates',
    'Current: Always have the latest data',
]
for pro in webhook_pros:
    doc.add_paragraph(f'✓ {pro}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph().add_run('POLLING (Pull Model)').bold = True
polling_cons = [
    'Delayed: Updates only when you check',
    'Inefficient: Constant checking even when nothing changes',
    'Resource intensive: Many unnecessary API calls',
    'Costly: Higher API usage and server costs',
    'Manual: Requires scheduling and maintenance',
    'Outdated: Data may be stale between checks',
]
for con in polling_cons:
    doc.add_paragraph(f'✗ {con}', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Example Scenario:').italic = True
doc.add_paragraph(
    'Polling: Your app checks Paystack every 5 minutes for new payments\n'
    '  • 288 API calls per day (24 hrs × 12 checks/hr)\n'
    '  • Most calls return "no new data"\n'
    '  • Up to 5-minute delay in notifications\n\n'
    'Webhooks: Paystack notifies you immediately\n'
    '  • Only calls when payment occurs\n'
    '  • Zero wasted API calls\n'
    '  • Instant notifications'
)

# Paystack Webhook Events
doc.add_heading('Paystack Webhook Events', level=2)
doc.add_paragraph()
doc.add_paragraph('Paystack sends webhooks for various events:')
doc.add_paragraph()

paystack_events = {
    'Payment Events': [
        'charge.success - Payment completed successfully',
        'charge.failed - Payment attempt failed',
    ],
    'Transfer Events': [
        'transfer.success - Transfer completed',
        'transfer.failed - Transfer failed',
        'transfer.reversed - Transfer reversed',
    ],
    'Subscription Events': [
        'subscription.create - New subscription created',
        'subscription.disable - Subscription cancelled',
        'subscription.not_renew - Subscription won\'t renew',
    ],
    'Customer Events': [
        'customeridentification.success - Customer verified',
        'customeridentification.failed - Verification failed',
    ],
    'Dispute Events': [
        'dispute.create - New dispute opened',
        'dispute.resolve - Dispute resolved',
    ],
}

for category, events in paystack_events.items():
    doc.add_heading(category, level=3)
    for event in events:
        doc.add_paragraph(event, style='List Bullet')

# Security
doc.add_page_break()
doc.add_heading('Webhook Security', level=2)
doc.add_paragraph()
doc.add_paragraph(
    'Security is critical for webhooks since they\'re public endpoints. Anyone could potentially send fake '
    'webhook requests to your server. Here\'s how to secure them:'
)

doc.add_heading('1. Signature Verification', level=3)
doc.add_paragraph('Paystack signs each webhook with your secret key using HMAC SHA512:')
doc.add_paragraph(
    '# Verify signature\n'
    'import hmac\n'
    'import hashlib\n\n'
    'def verify_signature(payload, signature, secret):\n'
    '    computed = hmac.new(\n'
    '        secret.encode(\'utf-8\'),\n'
    '        payload,\n'
    '        hashlib.sha512\n'
    '    ).hexdigest()\n'
    '    return computed == signature\n\n'
    '# In webhook handler\n'
    'signature = request.headers.get("x-paystack-signature")\n'
    'if not verify_signature(body, signature, SECRET_KEY):\n'
    '    return HTTPException(status_code=401, detail="Invalid signature")',
    style='Normal'
)

doc.add_heading('2. HTTPS Only', level=3)
doc.add_paragraph('Always use HTTPS URLs for webhooks to encrypt data in transit.')

doc.add_heading('3. Idempotency', level=3)
doc.add_paragraph(
    'Handle duplicate webhooks gracefully. Paystack may send the same webhook multiple times if it doesn\'t '
    'receive a response. Use transaction reference to detect duplicates.'
)

doc.add_heading('4. Validate Event Data', level=3)
doc.add_paragraph('Always validate that the event data makes sense before processing:')
validation_checks = [
    'Check that the transaction reference exists in your database',
    'Verify the amount matches your expected amount',
    'Confirm the event type is what you expect',
    'Validate the timestamp is recent',
]
for check in validation_checks:
    doc.add_paragraph(check, style='List Bullet')

doc.add_heading('5. Rate Limiting', level=3)
doc.add_paragraph('Protect against webhook flood attacks by implementing rate limiting.')

# Implementation
doc.add_heading('Webhook Implementation Example', level=2)
doc.add_paragraph()
doc.add_paragraph('Complete webhook handler with security and error handling:')
doc.add_paragraph(
    '@app.post("/webhook/paystack")\n'
    'async def paystack_webhook(request: Request):\n'
    '    try:\n'
    '        # Get signature from header\n'
    '        signature = request.headers.get("x-paystack-signature")\n'
    '        if not signature:\n'
    '            raise HTTPException(status_code=400, detail="No signature")\n'
    '        \n'
    '        # Get raw body\n'
    '        body = await request.body()\n'
    '        \n'
    '        # Verify signature\n'
    '        hash_value = hmac.new(\n'
    '            PAYSTACK_SECRET_KEY.encode(\'utf-8\'),\n'
    '            body,\n'
    '            hashlib.sha512\n'
    '        ).hexdigest()\n'
    '        \n'
    '        if hash_value != signature:\n'
    '            raise HTTPException(status_code=401, detail="Invalid signature")\n'
    '        \n'
    '        # Parse event\n'
    '        event = await request.json()\n'
    '        event_type = event.get("event")\n'
    '        data = event.get("data", {})\n'
    '        \n'
    '        # Handle different event types\n'
    '        if event_type == "charge.success":\n'
    '            reference = data.get("reference")\n'
    '            amount = data.get("amount") / 100\n'
    '            \n'
    '            # Check for duplicate\n'
    '            if already_processed(reference):\n'
    '                return {"status": "duplicate"}\n'
    '            \n'
    '            # Update database\n'
    '            update_transaction(reference, "success", amount)\n'
    '            \n'
    '            # Send confirmation email\n'
    '            send_email(data.get("customer").get("email"))\n'
    '            \n'
    '            # Mark as processed\n'
    '            mark_processed(reference)\n'
    '        \n'
    '        elif event_type == "charge.failed":\n'
    '            reference = data.get("reference")\n'
    '            update_transaction(reference, "failed")\n'
    '        \n'
    '        # Return success\n'
    '        return {"status": "success"}\n'
    '        \n'
    '    except Exception as e:\n'
    '        # Log error but return 200 to prevent retries\n'
    '        log_error(str(e))\n'
    '        return {"status": "error", "message": str(e)}',
    style='Normal'
)

# Benefits
doc.add_page_break()
doc.add_heading('Benefits of Webhooks', level=2)
doc.add_paragraph()

benefits = [
    'Real-time Updates: Get notified instantly when events occur, enabling immediate action',
    
    'Resource Efficiency: Eliminates the need for constant polling, reducing server load and API calls',
    
    'Scalability: Handles high volumes effortlessly - whether you have 10 or 10,000 transactions',
    
    'Cost Reduction: Fewer API calls mean lower costs and better rate limit management',
    
    'Better User Experience: Instant notifications lead to faster order processing and happier customers',
    
    'Reliability: Built-in retry mechanisms ensure you don\'t miss important events',
    
    'Automation: Enable fully automated workflows without manual intervention',
    
    'Flexibility: Subscribe only to events you care about',
    
    'Current Data: Always have the latest information without delays',
    
    'Reduced Complexity: Simpler architecture compared to polling systems',
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

# Use Cases
doc.add_heading('Common Webhook Use Cases', level=2)
doc.add_paragraph()

use_cases_webhooks = [
    'Payment Processing: Update order status when payment succeeds',
    'Email Notifications: Send receipts and confirmations automatically',
    'Inventory Management: Reduce stock when payment is confirmed',
    'Subscription Management: Handle subscription renewals and cancellations',
    'Fraud Detection: Get alerted immediately about suspicious transactions',
    'Accounting Integration: Sync transactions with accounting software',
    'Customer Support: Create support tickets for failed payments',
    'Analytics: Track payment success rates in real-time',
    'Refund Processing: Handle refund requests automatically',
    'Compliance: Maintain audit trails of all transaction events',
]
for use_case in use_cases_webhooks:
    doc.add_paragraph(use_case, style='List Bullet')

# Best Practices
doc.add_heading('Webhook Best Practices', level=2)
doc.add_paragraph()

webhook_best_practices = [
    'Always verify signatures - Never trust incoming webhooks without verification',
    
    'Respond quickly - Process webhooks asynchronously and return 200 OK immediately',
    
    'Handle idempotency - Use transaction IDs to detect and ignore duplicates',
    
    'Log everything - Keep detailed logs of all webhook deliveries for debugging',
    
    'Implement retries - Have your own retry logic for failed processing',
    
    'Use HTTPS - Never use HTTP for webhook endpoints',
    
    'Monitor webhook health - Track delivery rates and response times',
    
    'Test thoroughly - Use test events from Paystack dashboard',
    
    'Handle all event types gracefully - Don\'t crash on unexpected events',
    
    'Keep processing light - Move heavy processing to background jobs',
]
for practice in webhook_best_practices:
    doc.add_paragraph(practice, style='List Bullet')

# Summary
doc.add_heading('Summary', level=2)
doc.add_paragraph()
doc.add_paragraph(
    'Webhooks are a powerful, efficient, and modern way to build integrations. They enable real-time, '
    'event-driven architectures that are more scalable and cost-effective than traditional polling methods. '
    'For payment systems like Paystack, webhooks are essential for providing instant payment confirmations, '
    'enabling automated workflows, and delivering excellent user experiences. By implementing webhooks with '
    'proper security measures and best practices, you create robust, reliable applications that can handle '
    'payment processing at scale.'
)

# Save the document
doc.save('Section_2_Part2_Answers.docx')
print("Section 2 Part 2 answers document created successfully!")
